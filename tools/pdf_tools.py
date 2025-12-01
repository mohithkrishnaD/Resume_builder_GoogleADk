"""PDF Processing Tools - Pure ADK"""
# from google.adk.tools import tool
import PyPDF2
import docx
from typing import Dict
import re


# @tool
def parse_pdf_tool(file_path: str) -> Dict:

    
    """
    ADK Tool: Extract text from PDF file.
    Args:
    file_path: Path to PDF file
    Returns:
           Dictionary with extracted text and metadata
       """
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            return {
                "success": True,
                "text": text.strip(),
                "page_count": len(reader.pages)
            }
    except Exception as e:
        return {"success": False, "error": str(e), "text": ""}


# @tool
def parse_docx_tool(file_path: str) -> Dict:
    """
    ADK Tool: Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with extracted text
    """
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        return {
            "success": True,
            "text": text.strip(),
            "paragraph_count": len(doc.paragraphs)
        }
    except Exception as e:
        return {"success": False, "error": str(e), "text": ""}


# @tool
def extract_sections_tool(text: str) -> Dict:
    """
    ADK Tool: Extract common resume sections.

    Args:
        text: Full resume text

    Returns:
        Dictionary with extracted sections
    """
    sections = {
        "experience": "",
        "education": "",
        "skills": "",
        "summary": ""
    }

    text_lower = text.lower()

    # Experience
    exp_keywords = ["experience", "work history", "employment"]
    for kw in exp_keywords:
        if kw in text_lower:
            idx = text_lower.find(kw)
            sections["experience"] = text[idx:idx + 1500]
            break

    # Education
    edu_keywords = ["education", "academic", "degree"]
    for kw in edu_keywords:
        if kw in text_lower:
            idx = text_lower.find(kw)
            sections["education"] = text[idx:idx + 800]
            break

    # Skills
    skill_keywords = ["skills", "technical skills", "competencies"]
    for kw in skill_keywords:
        if kw in text_lower:
            idx = text_lower.find(kw)
            sections["skills"] = text[idx:idx + 500]
            break

    return sections


def extract_text_from_pdf(file_path: str) -> Dict:
    """
    Extract text from PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            return {
                "success": True,
                "text": text.strip(),
                "page_count": len(reader.pages),
                "word_count": len(text.split())
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "text": ""
        }


def extract_text_from_docx(file_path: str) -> Dict:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with extracted text
    """
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        return {
            "success": True,
            "text": text.strip(),
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(text.split())
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "text": ""
        }


def extract_resume_sections(text: str) -> Dict:
    """
    Extract common resume sections.

    Args:
        text: Full resume text

    Returns:
        Dictionary with extracted sections
    """
    sections = {
        "experience": "",
        "education": "",
        "skills": "",
        "summary": ""
    }

    text_lower = text.lower()

    # Extract Experience
    exp_keywords = ["experience", "work history", "employment", "professional experience"]
    for keyword in exp_keywords:
        if keyword in text_lower:
            idx = text_lower.find(keyword)
            sections["experience"] = text[idx:min(idx + 1500, len(text))]
            break

    # Extract Education
    edu_keywords = ["education", "academic", "qualification", "degree"]
    for keyword in edu_keywords:
        if keyword in text_lower:
            idx = text_lower.find(keyword)
            sections["education"] = text[idx:min(idx + 800, len(text))]
            break

    # Extract Skills
    skill_keywords = ["skills", "technical skills", "competencies", "expertise"]
    for keyword in skill_keywords:
        if keyword in text_lower:
            idx = text_lower.find(keyword)
            sections["skills"] = text[idx:min(idx + 500, len(text))]
            break

    return sections


def extract_contact_info(text: str) -> Dict:
    """
    Extract contact information from resume.

    Args:
        text: Resume text

    Returns:
        Dictionary with emails and phones
    """
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    phone_pattern = r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'

    emails = list(set(re.findall(email_pattern, text)))
    phones = list(set(re.findall(phone_pattern, text)))

    return {
        "emails": emails,
        "phones": phones
    }